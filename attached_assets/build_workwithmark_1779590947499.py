from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

sec = doc.sections[0]
sec.page_width  = int(8.5 * 914400)
sec.page_height = int(11  * 914400)
sec.left_margin = sec.right_margin = Inches(1)
sec.top_margin  = sec.bottom_margin = Inches(1)

NAVY  = RGBColor(0x1A, 0x3A, 0x5C)
OCEAN = RGBColor(0x00, 0x72, 0xBC)
RED   = RGBColor(0xC0, 0x39, 0x2B)
AMBER = RGBColor(0xE6, 0x7E, 0x00)
GREY  = RGBColor(0x55, 0x55, 0x55)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)
GOLD  = RGBColor(0xFF, 0xD7, 0x00)
GREEN = RGBColor(0x14, 0x5A, 0x32)
MONO  = RGBColor(0x2C, 0x2C, 0x2C)

def shd(para, fill):
    pPr = para._p.get_or_add_pPr()
    e = OxmlElement('w:shd')
    e.set(qn('w:val'), 'clear')
    e.set(qn('w:color'), 'auto')
    e.set(qn('w:fill'), fill)
    pPr.append(e)

def rule(colour="1A3A5C"):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    b = OxmlElement('w:bottom')
    b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '6')
    b.set(qn('w:space'), '1');    b.set(qn('w:color'), colour)
    pBdr.append(b); pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(8)

def section_bar(title, subtitle=""):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(4)
    shd(p, "1A3A5C")
    r1 = p.add_run(f"  {title}  ")
    r1.bold = True; r1.font.color.rgb = GOLD; r1.font.size = Pt(13)
    if subtitle:
        r2 = p.add_run(f"  —  {subtitle}")
        r2.font.color.rgb = RGBColor(0xAA,0xCC,0xEE)
        r2.font.size = Pt(11)

def body(text, colour=BLACK, size=11, indent=0, bold=False, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(6)
    if indent: p.paragraph_format.left_indent = Inches(indent)
    r = p.add_run(text)
    r.font.size = Pt(size); r.font.color.rgb = colour
    r.bold = bold; r.italic = italic
    return p

def code(text):
    p = doc.add_paragraph()
    shd(p, "F2F3F4")
    p.paragraph_format.left_indent  = Inches(0.2)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text)
    r.font.name = "Courier New"; r.font.size = Pt(9)
    r.font.color.rgb = MONO
    return p

def label(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text)
    r.bold = True; r.font.color.rgb = NAVY; r.font.size = Pt(11)

def note(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(4)
    shd(p, "EAF4FB")
    r = p.add_run(f"  NOTE:  {text}")
    r.italic = True; r.font.size = Pt(10); r.font.color.rgb = OCEAN

# ══════════════════════════════════════════════════════════════
# TITLE
# ══════════════════════════════════════════════════════════════
for txt, sz, col, fill in [
    ("STILL AFLOAT CRUISING",              22, WHITE, "1A3A5C"),
    ('"Work With Mark" — Implementation Brief', 14, GOLD, "1A3A5C"),
    ("For Replit Developer · Supabase Integration · Cloudflare Turnstile", 10,
     RGBColor(0xAA,0xCC,0xEE), "1A3A5C"),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shd(p, fill)
    r = p.add_run(txt); r.bold = True
    r.font.size = Pt(sz); r.font.color.rgb = col

doc.add_paragraph()

for lbl, val in [
    ("Page URL",      "/work-with-mark  (or /book)"),
    ("Nav label",     "Work With Mark"),
    ("DB",            "Supabase — new table: prospects"),
    ("Spam protection", "Cloudflare Turnstile (free, no limits)"),
    ("Future phase",  "TESS CRM connector — tess_synced / tess_id columns included now"),
]:
    p = doc.add_paragraph()
    r1 = p.add_run(f"{lbl}:  "); r1.bold = True
    r1.font.color.rgb = NAVY; r1.font.size = Pt(11)
    r2 = p.add_run(val); r2.font.size = Pt(11)

doc.add_paragraph(); rule()

# ══════════════════════════════════════════════════════════════
# SECTION 1 — PAGE COPY  (what visitors read)
# ══════════════════════════════════════════════════════════════
section_bar("SECTION 1", "Page Copy — What Visitors Read")

note("All text below is the live copy for the page. "
     "Formatting hierarchy: H1 → H2 → body paragraphs. "
     "Maintain the Still Afloat brand voice: direct, warm, no fluff.")

doc.add_paragraph()

# Hero
label("HERO — Above the fold")
body("Work With Mark", colour=NAVY, size=22, bold=True)
body("Plan your cruise with someone who's actually been there.",
     colour=OCEAN, size=14, italic=True)
body(
    "I'm Mark Millham — retired IT senior manager, U.S. veteran, former liveaboard sailor, "
    "and seven-cruise veteran. I'm an independent travel advisor affiliated with "
    "Cornerstone Collective, and I help people cut through the noise and book the right cruise "
    "for their situation — not the one that's easiest to sell.")

rule("0072BC")

# Why work with Mark
label("WHY WORK WITH MARK")
body("What you get that a booking engine doesn't give you", colour=NAVY, size=13, bold=True)
for bullet in [
    ("Someone who has cruised seven times across multiple lines, not someone reading from a brochure."),
    ("An IT-trained mind applied to cruise math — pricing, cabin categories, onboard costs, "
     "and whether a deal is actually a deal."),
    ("Real experience planning a cruise for a first-timer. My wife Ceci's first cruise "
     "was on the Norwegian Getaway. I know what makes a first cruise work — "
     "and what kills it."),
    ("A military background that means I take logistics seriously. "
     "Embarkation, port timing, travel days, contingency planning — "
     "these are not afterthoughts."),
    ("No pressure, no upsell agenda. I get paid when you sail happy, "
     "not when I push you into a cabin category you don't need."),
]:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(bullet); r.font.size = Pt(11)

doc.add_paragraph()
body(
    "I work with first-timers, returning cruisers planning something bigger, "
    "and couples where one partner is doing all the research. "
    "If that last one sounds familiar — I've been that partner. I know how it goes.")

rule("0072BC")

# What I help with
label("WHAT I HELP WITH")
body("From the first question to the gangway", colour=NAVY, size=13, bold=True)
for item in [
    ("Cruise line and ship selection",
     "Matching the right ship to your travel style, group size, and budget."),
    ("Itinerary planning",
     "Ports that are worth the day, and ports that look better on paper than in person."),
    ("Cabin category advice",
     "When the interior is fine, when the balcony changes everything, "
     "and what the upgrade actually costs you in real terms."),
    ("Understanding the real cost",
     "Gratuities, drink packages, specialty dining, excursions, Wi-Fi — "
     "I'll walk you through what's included and what adds up."),
    ("Casino and loyalty offers",
     "What cruise casino offers actually include, how to evaluate them, "
     "and whether chasing an offer makes sense for your situation."),
    ("Booking and documentation",
     "From deposit to final payment, passport requirements, travel insurance, "
     "and everything in between."),
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    r1 = p.add_run(item[0] + " — "); r1.bold = True
    r1.font.color.rgb = NAVY; r1.font.size = Pt(11)
    r2 = p.add_run(item[1]); r2.font.size = Pt(11)

rule("0072BC")

# Spanish block
label("SPANISH VERSION — Insert below English 'What I help with' section")
body(
    "¿Prefieres planear tu crucero en español? También trabajo con viajeros de habla hispana. "
    "Completa el formulario y cuéntame lo que necesitas.",
     colour=NAVY, size=11, italic=True)
note("Display this block on the Spanish version of the page. "
     "The contact form below is shared — language field will capture preference.")

rule("0072BC")

# Social proof placeholder
label("SOCIAL PROOF — Placeholder section (add testimonials when available)")
body("What cruisers are saying", colour=NAVY, size=13, bold=True)
body(
    "[PLACEHOLDER — insert first client testimonial here once available. "
    "Until then, this section can be omitted or replaced with a one-line trust statement "
    "such as: 'No booking fees. No pressure. Just someone who knows the water.']",
    colour=GREY, italic=True)

rule("0072BC")

# CTA above form
label("CALL TO ACTION — Immediately above the contact form")
body("Ready to start planning?", colour=NAVY, size=16, bold=True)
body(
    "Fill out the form below and I'll be in touch within 24 hours. "
    "No commitment, no sales pitch — just a conversation about where you want to go "
    "and how to get there without overpaying or overthinking it.")
body(
    "¿Prefieres comunicarte en español? Indícalo en el formulario.",
     colour=GREY, size=10, italic=True)

rule()

# ══════════════════════════════════════════════════════════════
# SECTION 2 — SUPABASE TABLE
# ══════════════════════════════════════════════════════════════
section_bar("SECTION 2", "Supabase — CREATE TABLE prospects")

note("Run this SQL in the Supabase SQL editor. "
     "tess_synced and tess_id are included now to support future TESS CRM connector.")

doc.add_paragraph()

sql_lines = [
    "-- ================================================",
    "-- Still Afloat LLC — prospects table",
    "-- Supabase SQL Editor",
    "-- ================================================",
    "",
    "CREATE TABLE IF NOT EXISTS prospects (",
    "  id               UUID        DEFAULT gen_random_uuid() PRIMARY KEY,",
    "  created_at       TIMESTAMPTZ DEFAULT NOW() NOT NULL,",
    "",
    "  -- Contact",
    "  first_name       TEXT        NOT NULL,",
    "  last_name        TEXT        NOT NULL,",
    "  email            TEXT        NOT NULL,",
    "  phone            TEXT,",
    "  preferred_lang   TEXT        DEFAULT 'en',   -- 'en' | 'es'",
    "",
    "  -- Trip details",
    "  num_travelers    INTEGER,",
    "  travel_dates     TEXT,           -- free text, e.g. 'March 2027'",
    "  destination      TEXT,           -- Caribbean | Mediterranean | Alaska | etc.",
    "  cruise_line_pref TEXT,           -- NCL | Royal | Carnival | No preference | etc.",
    "  budget_range     TEXT,           -- see form options below",
    "  first_time       BOOLEAN,",
    "  notes            TEXT,",
    "",
    "  -- Attribution",
    "  referral_source  TEXT,           -- YouTube | Website | Social | Referral | Other",
    "",
    "  -- CRM workflow",
    "  status           TEXT        DEFAULT 'new',  -- new | contacted | qualified | booked | lost",
    "  contacted_at     TIMESTAMPTZ,",
    "  advisor_notes    TEXT,",
    "",
    "  -- TESS connector (future phase)",
    "  tess_synced      BOOLEAN     DEFAULT FALSE,",
    "  tess_id          TEXT",
    ");",
    "",
    "-- Index for fast status queries",
    "CREATE INDEX idx_prospects_status    ON prospects(status);",
    "CREATE INDEX idx_prospects_email     ON prospects(email);",
    "CREATE INDEX idx_prospects_created   ON prospects(created_at DESC);",
    "",
    "-- Row Level Security (important — Supabase default is open)",
    "ALTER TABLE prospects ENABLE ROW LEVEL SECURITY;",
    "",
    "-- Only service_role key can read/write (form uses insert-only anon policy below)",
    "CREATE POLICY \"service_role_all\" ON prospects",
    "  FOR ALL USING (auth.role() = 'service_role');",
    "",
    "-- Anon INSERT only (form submission — no read access)",
    "CREATE POLICY \"anon_insert\" ON prospects",
    "  FOR INSERT WITH CHECK (true);",
]
for line in sql_lines:
    code(line)

rule()

# ══════════════════════════════════════════════════════════════
# SECTION 3 — CLOUDFLARE TURNSTILE SETUP
# ══════════════════════════════════════════════════════════════
section_bar("SECTION 3", "Cloudflare Turnstile — Setup Instructions")

body("Why Turnstile over reCAPTCHA", colour=NAVY, bold=True)
for point in [
    "Completely free — no usage tiers, no limits, no billing setup required.",
    "No privacy-invasive tracking. GDPR-compliant. Matters for Spanish/international visitors.",
    "Invisible to real users — a simple checkbox, no puzzles or image grids.",
    "Server-side token verification prevents form spoofing.",
    "Takes ~15 minutes to set up.",
]:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(point); r.font.size = Pt(10)

doc.add_paragraph()
label("STEP 1 — Create a Turnstile site")
body("1. Go to dash.cloudflare.com → Turnstile → Add site", indent=0.2)
body("2. Choose 'Managed' challenge type (recommended)", indent=0.2)
body("3. Enter your domain: stillafloatcruising.com", indent=0.2)
body("4. Copy the SITE KEY (public — goes in the HTML)", indent=0.2)
body("5. Copy the SECRET KEY (private — goes in Replit environment variables only)", indent=0.2)

note("Never put the SECRET KEY in client-side code or commit it to git. "
     "Store it as a Replit Secret: TURNSTILE_SECRET_KEY")

label("STEP 2 — Add to Replit Secrets")
for s in [
    "TURNSTILE_SECRET_KEY = <your secret key from Cloudflare>",
    "SUPABASE_URL         = <your Supabase project URL>",
    "SUPABASE_SERVICE_KEY = <your Supabase service_role key>",
]:
    code(f"  {s}")

note("Use service_role key on the server only — never expose it client-side. "
     "The anon key is safe client-side for the Supabase JS client.")

rule()

# ══════════════════════════════════════════════════════════════
# SECTION 4 — CONTACT FORM FIELDS SPEC
# ══════════════════════════════════════════════════════════════
section_bar("SECTION 4", "Contact Form — Fields & Validation Spec")

note("Map each field directly to the prospects table column of the same name.")

fields = [
    ("first_name",      "First Name",            "text",     "Required"),
    ("last_name",       "Last Name",             "text",     "Required"),
    ("email",           "Email Address",         "email",    "Required - validate format"),
    ("phone",           "Phone Number",          "tel",      "Optional - US and international formats"),
    ("preferred_lang",  "Preferred Language",    "select",   "Options: English / Espanol (default: English)"),
    ("num_travelers",   "Number of Travelers",   "number",   "Required - min 1, max 20"),
    ("travel_dates",    "Approximate Travel Dates", "text",  "Required - placeholder: e.g. Spring 2027"),
    ("destination",     "Destination Interest",  "select",   "Caribbean, Mediterranean, Alaska, Bahamas, Mexico, Europe, Not sure yet"),
    ("cruise_line_pref","Cruise Line Preference", "select",  "Norwegian NCL, Royal Caribbean, Carnival, Celebrity, Holland America, MSC, No preference"),
    ("budget_range",    "Budget Per Person",     "select",   "Under $1000, $1000-$2000, $2000-$3500, $3500-$5000, $5000+"),
    ("first_time",      "Is this your first cruise?", "radio", "Yes / No"),
    ("referral_source", "How did you find us?",  "select",   "YouTube, Website, Social Media, Friend or family referral, Other"),
    ("notes",           "Questions or anything else we should know", "textarea", "Optional - 4 rows, max 1000 characters"),
    ("cf-turnstile",    "Cloudflare Turnstile widget", "widget", "Required - verify before submit"),
]

for db_col, label_txt, ftype, notes_txt in fields:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(1)
    r1 = p.add_run(f"{db_col:<22}")
    r1.font.name = "Courier New"; r1.font.size = Pt(9.5)
    r1.font.color.rgb = OCEAN; r1.bold = True
    r2 = p.add_run(f"  {label_txt}")
    r2.font.size = Pt(10.5); r2.font.color.rgb = BLACK; r2.bold = True
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Inches(0.3)
    p2.paragraph_format.space_after = Pt(4)
    r3 = p2.add_run(f"type: {ftype}   |   {notes_txt}")
    r3.font.size = Pt(9.5); r3.font.color.rgb = GREY

rule()

# ══════════════════════════════════════════════════════════════
# SECTION 5 — SERVER ENDPOINT SPEC
# ══════════════════════════════════════════════════════════════
section_bar("SECTION 5", "Backend Endpoint Spec — /api/contact")

note("The form POSTs JSON to a Replit backend endpoint (Node/Express or Python/Flask). "
     "The endpoint does TWO things in order: "
     "(1) verify the Turnstile token with Cloudflare, "
     "(2) insert the prospect row into Supabase. "
     "If Turnstile verification fails, do NOT insert and return 400.")

doc.add_paragraph()
label("Request body (JSON)")
for line in [
    "{",
    '  "first_name":      "Jane",',
    '  "last_name":       "Smith",',
    '  "email":           "jane@example.com",',
    '  "phone":           "+1-555-123-4567",',
    '  "preferred_lang":  "en",',
    '  "num_travelers":   2,',
    '  "travel_dates":    "March 2027",',
    '  "destination":     "Caribbean",',
    '  "cruise_line_pref":"No preference",',
    '  "budget_range":    "$2,000–$3,500",',
    '  "first_time":      true,',
    '  "referral_source": "YouTube",',
    '  "notes":           "Looking for a 7-night western Caribbean.",',
    '  "cf-turnstile-response": "<token from widget>"',
    "}",
]:
    code(line)

doc.add_paragraph()
label("Pseudocode — endpoint logic")
for line in [
    "POST /api/contact",
    "",
    "1. Parse JSON body",
    "2. POST to https://challenges.cloudflare.com/turnstile/v0/siteverify",
    "   body: { secret: TURNSTILE_SECRET_KEY, response: body['cf-turnstile-response'] }",
    "3. If siteverify.success !== true → return 400 { error: 'Verification failed' }",
    "4. Build prospect object from validated fields",
    "5. supabase.from('prospects').insert(prospect_object)",
    "6. If Supabase error → return 500 { error: 'Could not save your request' }",
    "7. Return 200 { message: 'Thank you — Mark will be in touch within 24 hours.' }",
]:
    code(line)

note("Send a confirmation email to the prospect's email address on success "
     "if Replit has email capability (Resend, SendGrid, or Nodemailer). "
     "Even a simple auto-reply builds trust immediately.")

rule()

# ══════════════════════════════════════════════════════════════
# SECTION 6 — SUCCESS / ERROR STATES
# ══════════════════════════════════════════════════════════════
section_bar("SECTION 6", "Form UX — Success & Error Copy")

label("Success message (replace form on 200 response)")
body(
    "You're on the radar, [First Name].",
    colour=NAVY, size=14, bold=True)
body(
    "Mark will review your details and be in touch within 24 hours. "
    "In the meantime, catch up on the channel — "
    "there's already something there worth watching.",
    colour=BLACK, size=11)

label("Error — Turnstile failed")
body(
    "We couldn't verify that you're human — it happens. "
    "Please refresh the page and try again.",
    colour=RED, size=11)

label("Error — Server/Supabase failure")
body(
    "Something went wrong on our end. "
    "Please email Mark directly at [EMAIL ADDRESS] and he'll get back to you.",
    colour=RED, size=11)

label("Validation errors (inline, per field)")
body(
    "Keep these brief and specific: "
    "'Please enter a valid email address.' "
    "'Please tell us how many travelers.' "
    "Avoid generic 'This field is required' messages.",
    colour=GREY, size=10, italic=True)

rule()

# ══════════════════════════════════════════════════════════════
# SECTION 7 — TESS CONNECTOR (FUTURE PHASE)
# ══════════════════════════════════════════════════════════════
section_bar("SECTION 7", "TESS CRM Connector — Future Phase Notes")

body(
    "The prospects table is pre-built for TESS sync. "
    "When Cornerstone Collective provides API access to TESS, the connector "
    "should follow this pattern:", colour=NAVY)

for step in [
    "Query Supabase for rows WHERE tess_synced = FALSE AND status != 'lost'",
    "For each row: POST or PUT to TESS API to create/update the prospect record",
    "On success: UPDATE prospects SET tess_synced = TRUE, tess_id = '<TESS record ID>'",
    "On failure: log the error, leave tess_synced = FALSE for retry",
    "Run on a schedule (cron) — every 15–30 minutes is sufficient for a solo advisor",
]:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(step); r.font.size = Pt(10)

doc.add_paragraph()
body(
    "Action item: Contact Cornerstone Collective and ask specifically for: "
    "(1) TESS API documentation, (2) whether they provide API keys for affiliated advisors, "
    "(3) whether webhook-based sync is an option. "
    "Bring that documentation and I can help design the connector.",
    colour=GREY, italic=True, size=10)

rule()

out = "/sessions/beautiful-cool-babbage/mnt/outputs/Still_Afloat_WorkWithMark_Brief.docx"
doc.save(out)
print(f"Saved: {out}")
